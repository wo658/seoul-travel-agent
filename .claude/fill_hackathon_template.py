#!/usr/bin/env python3
"""
Fill 2025 SeSAC Hackathon AI Service Proposal Template
Preserves original layout and fills with Seoul Travel Agent content
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def fill_template():
    # Load template
    doc = Document('/Users/jhkim/seoul-travel-agent/2025년 새싹 해커톤 AI 서비스 기획서 양식.docx')

    # Table 2: Team Information
    team_table = doc.tables[1]
    team_table.rows[0].cells[1].text = 'Seoul Travel Agent'
    team_table.rows[1].cells[1].text = '팀장: [팀장명], 팀원: [팀원명]'

    # Table 3: AI Service Name
    service_name_table = doc.tables[2]
    service_name_table.rows[0].cells[0].text = '''ㅇ 제안한 AI 서비스의 명칭

Seoul Travel Agent (서울 여행 AI 도우미)
- 영문명: Seoul Travel Agent - AI-Powered Travel Planning Assistant
- 한글명: 서울 여행 AI 도우미
- 부제: 사전 계획부터 실시간 가이드까지, 당신의 완벽한 서울 여행 파트너'''

    # Table 4: Data Sources
    data_table = doc.tables[3]

    # Row 1
    data_table.rows[1].cells[0].text = '1'
    data_table.rows[1].cells[1].text = '전국 관광지 정보 표준 데이터\nNaver 지역 검색 API\n서울 실시간 Open API (향후 활용 예정)'
    data_table.rows[1].cells[2].text = '관광/여행'
    data_table.rows[1].cells[3].text = '한국관광공사 TourAPI 4.0\nNaver Cloud Platform\n서울열린데이터광장'

    # Row 2
    data_table.rows[2].cells[0].text = '2'
    data_table.rows[2].cells[1].text = 'GPT-5-nano\nSentence Transformers (임베딩 모델)'
    data_table.rows[2].cells[2].text = 'AI 모델'
    data_table.rows[2].cells[3].text = 'OpenAI\nHugging Face'

    # Table 5: Core Content (핵심내용)
    core_content_table = doc.tables[4]
    core_content = '''ㅇ 서비스 아이디어의 핵심 내용

Seoul Travel Agent는 **두 가지 전문 AI 에이전트**로 구성된 서울 여행 종합 플래너입니다.

【1. 여행 계획 사전 수립 + 수정 (Planner Agent)】

▶ 사전 계획 수립 프로세스
  • 사용자 입력 수집: "3일간 서울 여행, 예산 50만원, 역사 문화 관심"
  • 선호도 분석: AI가 여행 기간, 예산, 관심사, 동행인 정보 파악
  • 관광지 검색: Vector DB에서 유사 관광지 자동 검색 (RAG 기술 활용)
  • 맞춤형 일정 생성: 시간대별 관광지, 이동 경로, 식당 추천 포함

▶ 계획 수정 기능
  • 실시간 피드백 반영: "2일차 점심을 한식으로 바꿔줘"
  • 부분 수정 지원: 전체 일정 유지하면서 특정 부분만 변경
  • 예산 재조정: 수정 후에도 예산 제약 자동 유지
  • 반복 수정 가능: 만족할 때까지 무제한 수정

▶ 기술적 차별점
  • Multi-Agent 시스템: LangGraph 기반 워크플로우 자동화
  • RAG (검색 증강 생성): AI 환각 방지, 실제 데이터 기반 추천
  • 실시간 정보 통합: Naver Places API로 영업시간, 휴무일 확인


【2. 실제 여행 도중 상황 대응 가이드 (Reviewer Agent)】

▶ 여행 중 발생 가능한 상황별 대응

1️⃣ 날씨 변화 대응
  • 상황: "비가 와서 야외 활동이 어려워요"
  • AI 대응: 실내 활동 중심으로 자동 재구성 (박물관, 카페, 쇼핑몰)
  • 기술: 날씨 API 연동 + 실내 장소 우선 검색

2️⃣ 체력 관리 및 일정 조정
  • 상황: "너무 피곤해서 오후 일정을 줄이고 싶어요"
  • AI 대응: 이동 거리 최소화, 휴식 시간 추가, 필수 코스만 유지
  • 기술: 피로도 분석 알고리즘 + 최적 경로 재계산

3️⃣ 돌발 이벤트 발생
  • 상황: "친구가 갑자기 합류했어요", "공연 티켓을 구했어요"
  • AI 대응: 인원수 변경 반영, 일정에 공연 시간 삽입
  • 기술: 동적 일정 재조정 + 시간 충돌 방지 알고리즘

4️⃣ 장소 휴무/혼잡도 대응
  • 상황: "목표했던 레스토랑이 휴무예요"
  • AI 대응: 인근 대안 장소 즉시 추천 (유사 카테고리, 예산 고려)
  • 기술: Naver API 실시간 조회 + Vector DB 유사도 검색

5️⃣ 예산 초과 방지
  • 상황: "예산이 부족해요"
  • AI 대응: 무료/저렴한 대안 제시, 남은 일정 예산 재분배
  • 기술: 실시간 예산 추적 + 최적화 알고리즘

▶ Reviewer Agent 워크플로우
  1. 피드백 분석: 사용자 의도 파악 (approve/reject/modify)
  2. 수정 범위 결정: "2일차 저녁만", "전체 일정" 등 자동 판단
  3. 맥락 유지 재계획: 기존 일정 흐름 보존하며 변경
  4. 검증 및 제시: 시간 중복, 이동 거리 재확인 후 제안

▶ 실시간 가이드 핵심 기능
  • 다중 반복 수정: 여행 내내 무제한 수정 가능
  • 맥락 보존: 이전 대화 기억 + 일정 연속성 유지
  • 즉시 대응: 평균 응답 시간 5-10초 (스트리밍)
  • 학습 개선: 사용자 선택 패턴 학습으로 추천 정확도 향상'''

    core_content_table.rows[0].cells[0].text = core_content

    # Table 6: Background and Purpose
    background_table = doc.tables[5]
    background = '''ㅇ AI 서비스를 구상하게 된 배경

【문제 인식】
1. 기존 여행 앱의 한계
  • 정적인 추천: 사용자 맥락 무시한 일반적 추천
  • 실시간 대응 부족: 여행 중 변경사항 대응 어려움
  • 데이터 신뢰성: 오래된 정보로 인한 실망 경험

2. 서울 여행자의 실제 니즈
  • 외국인 관광객: 언어 장벽, 복잡한 대중교통 시스템
  • 국내 관광객: 숨은 명소 발굴, 혼잡도 회피 욕구
  • 모든 여행자: 실시간 상황 변화 대응 필요

【서비스 목적】
1. 개인화된 여행 경험 제공
  • 사용자 선호도, 체력, 예산 맞춤형 일정
  • Multi-Agent 시스템으로 계획 수립(Planner) + 실시간 대응(Reviewer) 분리

2. 실시간 적응형 서비스
  • 여행 전: 완벽한 사전 계획 수립
  • 여행 중: 날씨, 교통, 컨디션 등 변수 즉시 반영

3. 신뢰할 수 있는 정보 제공
  • 공공 데이터 + 실시간 API 결합
  • RAG 시스템으로 AI 환각(Hallucination) 방지'''

    background_table.rows[0].cells[0].text = background

    # Table 7: Details
    details_table = doc.tables[6]
    details = '''ㅇ 활용 데이터 및 AI 모델

【데이터 소스】
1. 공공 오픈 데이터
  • 전국 관광지 정보 표준 데이터 (한국관광공사 TourAPI 4.0)
  • Naver 지역 검색 API (Naver Cloud Platform)
  • 서울 실시간 Open API (향후 활용 예정 - 서울열린데이터광장)

2. Vector Database (RAG)
  • ChromaDB: 관광지 설명 임베딩 저장
  • Sentence Transformers: 의미론적 유사도 검색

3. 실시간 데이터
  • Naver 지역 검색 API: 영업시간, 리뷰, 평점, 전화번호
  • 서울 실시간 교통/날씨 정보 (향후 연동 예정)

【AI 모델】
  • LLM: OpenAI GPT-5-nano
  • 임베딩: Sentence Transformers (all-MiniLM-L6-v2)
  • 워크플로우: LangGraph (Multi-Agent Orchestration)


ㅇ 세부내용 - 기술 아키텍처

【Frontend】
  • React Native 0.81 + Expo 54
  • NativeWind (Tailwind CSS for React Native)
  • TypeScript 기반 타입 안정성

【Backend】
  • FastAPI (Python 3.12+)
  • SQLAlchemy 2.0 (ORM)
  • PostgreSQL (프로덕션)

【AI Infrastructure】
  • LangGraph StateGraph: Multi-Agent 워크플로우
  • ChromaDB: 벡터 검색 엔진
  • SSE (Server-Sent Events): 실시간 스트리밍


ㅇ 서비스 방법 - 사용자 여정

【Step 1: 계획 생성 (Planner Agent)】
  입력: "3일간 서울 여행, 예산 50만원, 역사 문화 관심"
    ↓
  [정보 수집] → [관광지 검색] → [일정 생성] → [계획 제시]
    ↓
  출력: 일자별 상세 일정 (시간대별 관광지, 이동 경로, 식당 추천)

【Step 2: 계획 수정 (사전)】
  피드백: "2일차 점심을 한식으로 변경해주세요"
    ↓
  [피드백 분석] → [부분 수정] → [검증] → [제시]
    ↓
  출력: 수정된 일정 (기존 흐름 유지)

【Step 3: 여행 중 실시간 가이드 (Reviewer Agent)】
  상황: "비가 와서 야외 활동이 어려워요"
    ↓
  [상황 분석] → [대안 검색] → [일정 재구성] → [제시]
    ↓
  출력: 실내 활동 중심 수정 일정 (박물관, 카페, 쇼핑몰)


ㅇ 서비스 창의성 및 구현 가능성

【창의성】
✓ Multi-Agent 시스템: 계획(Planner) + 가이드(Reviewer) 역할 분리
✓ 양방향 대화형 UI: 채팅 기반 자연스러운 상호작용
✓ RAG 기술: Vector DB로 환각 없는 정확한 추천
✓ 실시간 적응: 여행 중 상황 변화 즉시 반영

【구현 가능성】
✓ MVP 완성: Planner/Reviewer Agent 구현 완료
✓ 검증된 기술: LangGraph, ChromaDB, FastAPI 안정적
✓ API 연동: Naver Places API 실제 통합 완료
✓ 크로스 플랫폼: React Native로 iOS/Android 동시 지원


ㅇ 서비스의 예상 UI/UX 이미지 시각화

【화면 1: 홈 화면】
┌─────────────────────────────────┐
│  Seoul Travel Agent             │
│  ───────────────────────────    │
│  💬 새로운 여행 계획 시작       │
│                                  │
│  📝 저장된 계획                 │
│  ├ 서울 3일 역사 문화 여행      │
│  └ 데이트 코스 1일 플랜         │
│                                  │
│  ⚡ 빠른 시작                   │
│  • 1일 코스  • 3일 코스         │
└─────────────────────────────────┘

【화면 2: 채팅 화면】
┌─────────────────────────────────┐
│  ← 새 여행 계획                 │
│  ───────────────────────────    │
│  You:                            │
│  3일간 서울 여행, 예산 50만원,  │
│  역사 문화 관심있어요           │
│                                  │
│  AI: 🤖 타이핑 중...            │
│  알겠습니다! 3일 일정을          │
│  생성하고 있어요...              │
│                                  │
│  [일정 미리보기 카드]            │
│  ┌────────────────────────┐    │
│  │ Day 1: 경복궁 → 광장시장│    │
│  │ Day 2: 북촌 → 인사동    │    │
│  └────────────────────────┘    │
│                                  │
│  ┌─────────────────────────┐   │
│  │ 메시지 입력...           │   │
│  └─────────────────────────┘   │
└─────────────────────────────────┘

【화면 3: 일정 뷰어 (타임라인)】
┌─────────────────────────────────┐
│  서울 3일 역사 문화 여행         │
│  2024.12.20 - 12.22             │
│  ───────────────────────────    │
│  [Day 1] 2024.12.20 (금)        │
│  ┌───────────────────────────┐ │
│  │ 09:00 - 11:00  경복궁     │ │
│  │ 🏛️ 관광지 | 3,000원      │ │
│  │ 📍 서울시 종로구...       │ │
│  │ [지도 보기] [수정 요청]   │ │
│  └───────────────────────────┘ │
│  ┌───────────────────────────┐ │
│  │ 12:00 - 13:30  광장시장   │ │
│  │ 🍜 음식점 | 15,000원      │ │
│  └───────────────────────────┘ │
│                                  │
│  [Day 2] 2024.12.21 (토) ...    │
│  [Day 3] 2024.12.22 (일) ...    │
│                                  │
│  💾 저장  📤 공유  ✏️ 전체 수정 │
└─────────────────────────────────┘'''

    details_table.rows[0].cells[0].text = details

    # Table 8: Expected Effects
    effects_table = doc.tables[7]
    effects = '''ㅇ 사회/경제적 파급(기대) 효과

【사용자 측면】
1. 시간 절감
  • 기존: 여행 계획 수립 5-10시간 → 개선: 10-15분 (90% 단축)

2. 개인화 경험
  • 획일적 패키지 투어 탈피
  • 개인 선호도, 체력, 예산 맞춤형 일정

3. 실시간 대응력
  • 날씨 변화, 돌발 상황 즉시 대응
  • 피로도, 컨디션 고려한 일정 조정

【비즈니스 측면】
1. 서울 관광 산업 활성화
  • 연간 서울 방문 외국인: 약 1,300만 명 (2023년)
  • 잠재 사용자: 국내외 관광객 연간 2,000만 명
  • 비주류 관광지 홍보로 골목 상권 활성화

2. 수익 모델
  • 프리미엄 기능: 무제한 일정 수정, 우선 응답
  • 제휴 수수료: 숙박, 식당, 투어 예약 연계
  • 광고 수익: 관광지, 지역 상권 프로모션

3. 확장 가능성
  • 다른 도시: 부산, 제주, 경주
  • B2B 서비스: 여행사, 호텔 제휴
  • 해외 진출: K-문화 관심 증가

【사회적 측면】
1. 접근성 향상
  • 다국어 지원 (한/영/중/일)
  • 장애인 관광 지원 (휠체어 접근 가능 코스)
  • 고령층 친화적 UI

2. 지역 경제 활성화
  • 비수기 관광 수요 창출 (계절별 테마 코스)
  • 신규 관광 콘텐츠 발굴

3. 환경 기여
  • 효율적 경로로 탄소 배출 감소
  • 대중교통 우선 추천'''

    effects_table.rows[0].cells[0].text = effects

    # Save filled document
    output_path = '/Users/jhkim/seoul-travel-agent/2025년_새싹_해커톤_Seoul_Travel_Agent_기획서.docx'
    doc.save(output_path)
    print(f'✅ Document saved: {output_path}')
    print(f'📄 File size: {len(open(output_path, "rb").read()) / 1024:.1f} KB')

if __name__ == '__main__':
    fill_template()
