#!/usr/bin/env python3
"""
Markdown to Word Document Converter
Converts the Seoul Travel AI Agent proposal from Markdown to Word format
"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_styled_document():
    """Create a Word document with custom styles"""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(11)

    # Title style
    title_style = doc.styles['Title']
    title_style.font.name = '맑은 고딕'
    title_style.font.size = Pt(20)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 51, 102)

    # Heading styles
    for i in range(1, 5):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = '맑은 고딕'
        heading_style.font.color.rgb = RGBColor(0, 51, 102)
        heading_style.font.size = Pt(18 - (i * 2))
        heading_style.font.bold = True

    return doc

def parse_markdown_line(line):
    """Parse a markdown line and return its type and content"""
    line = line.rstrip()

    # Heading
    if line.startswith('#'):
        level = len(re.match(r'^#+', line).group())
        content = line.lstrip('#').strip()
        return 'heading', level, content

    # Horizontal rule
    if line.strip() in ['---', '***', '___']:
        return 'hr', None, None

    # Code block
    if line.startswith('```'):
        return 'code_fence', None, line.strip('`').strip()

    # Bullet list
    if re.match(r'^\s*[-*]\s+', line):
        indent = len(re.match(r'^\s*', line).group())
        content = re.sub(r'^\s*[-*]\s+', '', line)
        return 'bullet', indent, content

    # Empty line
    if not line.strip():
        return 'empty', None, None

    # Regular paragraph
    return 'paragraph', None, line

def add_formatted_text(paragraph, text):
    """Add text to paragraph with inline formatting (bold, italic, code)"""
    # Handle bold
    parts = re.split(r'(\*\*.*?\*\*)', text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part.strip('*'))
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part.strip('`'))
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)

def convert_md_to_docx(md_file_path, docx_file_path):
    """Convert Markdown file to Word document"""
    doc = create_styled_document()

    with open(md_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_content = []
    code_language = None

    i = 0
    while i < len(lines):
        line = lines[i]
        line_type, level_or_indent, content = parse_markdown_line(line)

        # Handle code blocks
        if line_type == 'code_fence':
            if not in_code_block:
                in_code_block = True
                code_language = content
                code_content = []
            else:
                # End of code block
                in_code_block = False
                if code_content:
                    code_para = doc.add_paragraph()
                    code_para.style = 'Normal'
                    code_run = code_para.add_run('\n'.join(code_content))
                    code_run.font.name = 'Courier New'
                    code_run.font.size = Pt(9)
                    code_para.paragraph_format.left_indent = Inches(0.5)
                    code_para.paragraph_format.space_before = Pt(6)
                    code_para.paragraph_format.space_after = Pt(6)
                code_content = []
            i += 1
            continue

        if in_code_block:
            code_content.append(line.rstrip())
            i += 1
            continue

        # Handle headings
        if line_type == 'heading':
            if level_or_indent == 1:
                para = doc.add_heading(content, level=0)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_heading(content, level=level_or_indent)

        # Horizontal rule
        elif line_type == 'hr':
            doc.add_paragraph()  # Add spacing

        # Bullet list
        elif line_type == 'bullet':
            para = doc.add_paragraph(style='List Bullet')
            add_formatted_text(para, content)

            # Handle indentation
            if level_or_indent > 0:
                para.paragraph_format.left_indent = Inches(level_or_indent / 4)

        # Empty line
        elif line_type == 'empty':
            pass  # Skip empty lines

        # Regular paragraph
        elif line_type == 'paragraph':
            para = doc.add_paragraph()
            add_formatted_text(para, content)

        i += 1

    # Save document
    doc.save(docx_file_path)
    print(f"✅ Word document created: {docx_file_path}")

if __name__ == "__main__":
    md_file = "/Users/jhkim/seoul-travel-agent/서울여행AI에이전트_기획서.md"
    docx_file = "/Users/jhkim/seoul-travel-agent/서울여행AI에이전트_기획서.docx"

    convert_md_to_docx(md_file, docx_file)
